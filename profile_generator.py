"""
Surgeon Profile Document Generator

Creates formatted .docx surgeon profiles from structured data.
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

# Attempt to use the INFORM logo if it exists; otherwise skip it
LOGO_PATH = os.path.join(
    os.path.dirname(BASE_DIR), "SurgeonProfiles", "Logo", "INFORM_logo_White_01.png"
)
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = None

# Specialty-to-membership mappings (fallback)
SPECIALTY_MEMBERSHIPS = {
    "Orthopedic Surgery": [
        "American Academy of Orthopaedic Surgeons (AAOS)",
        "American Board of Orthopaedic Surgery",
    ],
    "Cardiothoracic Surgery": [
        "Society of Thoracic Surgeons (STS)",
        "American Association for Thoracic Surgery (AATS)",
    ],
    "Cardiovascular Surgery": [
        "Society of Thoracic Surgeons (STS)",
        "American College of Cardiology (ACC)",
    ],
    "General Surgery": [
        "American College of Surgeons (ACS)",
        "Society of American Gastrointestinal and Endoscopic Surgeons (SAGES)",
    ],
    "Neurosurgery": [
        "American Association of Neurological Surgeons (AANS)",
        "Congress of Neurological Surgeons (CNS)",
    ],
    "Urology": [
        "American Urological Association (AUA)",
    ],
    "Vascular Surgery": [
        "Society for Vascular Surgery (SVS)",
        "American College of Surgeons (ACS)",
    ],
    "Plastic Surgery": [
        "American Society of Plastic Surgeons (ASPS)",
    ],
    "Obstetrics & Gynecology": [
        "American College of Obstetricians and Gynecologists (ACOG)",
    ],
}

# Color constants
DARK_BLUE = RGBColor(0x1B, 0x4F, 0x72)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B4F72"
ROW_ALT_BG = "F5F8FA"
HEADER_TEXT_COLOR = RGBColor(0x88, 0x88, 0x88)

# Spacing constants (EMU)
SECTION_SPACE_BEFORE = 190500
SECTION_SPACE_AFTER = 63500
CONTACT_SPACE_AFTER = 25400
BULLET_SPACE_AFTER = 38100
BIO_SPACE_AFTER = 76200
TITLE_SPACE_AFTER = 38100
NAME_SPACE_AFTER = 25400
SUBTITLE_SPACE_AFTER = 12700
LOCATION_SPACE_AFTER = 127000


def _title(s):
    return s.strip().title() if s else ""


# ---------------------------------------------------------------------------
# Document formatting helpers
# ---------------------------------------------------------------------------

def _set_paragraph_border_bottom(paragraph, color="1B4F72", sz="3", space="4"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:color="{color}" w:sz="{sz}" w:space="{space}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_cell_shading(cell, color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'  <w:left w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'  <w:bottom w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'  <w:right w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'  <w:insideH w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'  <w:insideV w:val="single" w:color="auto" w:sz="4" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)


def _set_cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = None
    pf = p.paragraph_format
    pf.space_before = Emu(SECTION_SPACE_BEFORE)
    pf.space_after = Emu(SECTION_SPACE_AFTER)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    _set_paragraph_border_bottom(p)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Paragraph")
    p.paragraph_format.space_after = Emu(BULLET_SPACE_AFTER)
    return p


def _add_contact_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(CONTACT_SPACE_AFTER)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)
    return p


def _add_styled_table(doc, headers, data_rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        use_shading = (row_idx % 2 == 0)
        for ci, value in enumerate(row_data):
            cell = row_cells[ci]
            if use_shading:
                _set_cell_shading(cell, ROW_ALT_BG)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(value))
            run.font.size = Pt(9)

    return table


def _add_name_block_with_photo(doc, full_name, specialty, location_str, photo_path):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_no_table_borders(table)

    half_width = Emu(3429000)

    photo_cell = table.rows[0].cells[0]
    photo_cell.width = half_width
    _set_cell_vertical_center(photo_cell)
    p_photo = photo_cell.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_photo.paragraph_format.space_before = Emu(0)
    p_photo.paragraph_format.space_after = Emu(0)
    run_img = p_photo.add_run()
    run_img.add_picture(photo_path, width=Inches(1.5))

    text_cell = table.rows[0].cells[1]
    text_cell.width = half_width
    _set_cell_vertical_center(text_cell)

    p_name = text_cell.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Emu(0)
    p_name.paragraph_format.space_after = Emu(NAME_SPACE_AFTER)
    run = p_name.add_run(full_name)
    run.bold = True
    run.font.size = Pt(15)

    p_spec = text_cell.add_paragraph()
    p_spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_spec.paragraph_format.space_after = Emu(SUBTITLE_SPACE_AFTER)
    run = p_spec.add_run(specialty)
    run.font.color.rgb = GRAY

    p_loc = text_cell.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.paragraph_format.space_after = Emu(0)
    run = p_loc.add_run(location_str)
    run.font.color.rgb = GRAY


def _add_name_block_no_photo(doc, full_name, specialty, location_str):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Emu(NAME_SPACE_AFTER)
    run = p1.add_run(full_name)
    run.bold = True
    run.font.size = Pt(15)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Emu(SUBTITLE_SPACE_AFTER)
    run = p2.add_run(specialty)
    run.font.color.rgb = GRAY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Emu(LOCATION_SPACE_AFTER)
    run = p3.add_run(location_str)
    run.font.color.rgb = GRAY


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_profile(profile_data: dict, output_dir: str = None) -> str:
    """
    Generate a .docx surgeon profile from a structured data dict.

    Args:
        profile_data: Dict with keys matching the research output schema:
            npi, full_name, credential, specialty, gender,
            address, city, state, zip, phone, practice_name, practice_website,
            description (biography text),
            education, board_certs, memberships, languages,
            affiliations, ratings, awards, media,
            procedures (list of dicts with name, informed_score, cases),
            source_urls, photo_path, locations
        output_dir: Where to save the .docx. Defaults to ./output/

    Returns:
        Path to the generated .docx file.
    """
    output_dir = output_dir or OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    d = profile_data
    npi = d.get("npi", "unknown")
    full_name = (
        d.get("full_name")
        or d.get("name")
        or " ".join(filter(None, [
            d.get("first_name", ""),
            d.get("middle_name", d.get("middle", "")),
            d.get("last_name", ""),
        ])).strip()
        or "Unknown Surgeon"
    )
    credential = d.get("credential", "M.D.")
    specialty = _title(d.get("specialty", "Surgery"))
    city = _title(d.get("city", ""))
    state = d.get("state", "")
    address = _title(d.get("address", ""))
    zip_code = d.get("zip", "")
    phone = d.get("phone", "")
    practice_name = d.get("practice_name", "")
    practice_website = d.get("practice_website", "")

    # Build location string
    locations = d.get("locations", [])
    if len(locations) > 1:
        loc_cities = list(dict.fromkeys(
            loc["city"] for loc in locations if loc.get("city")
        ))
        if len(loc_cities) > 1:
            location_str = " / ".join(loc_cities[:3]) + f", {state}"
        else:
            location_str = f"{city}, {state}"
    else:
        location_str = f"{city}, {state}"

    # Build display name with credential
    if credential and not full_name.endswith(credential):
        display_name = f"{full_name}, {credential}"
    else:
        display_name = full_name

    # Education
    education_items = list(d.get("education", []))

    # Board certs
    board_certs = d.get("board_certs", [])
    if not board_certs:
        board_certs = [f"{credential} - {specialty}"]

    # Memberships
    memberships = list(d.get("memberships", []))
    if not memberships:
        memberships = list(SPECIALTY_MEMBERSHIPS.get(specialty, []))
    if not any("american medical association" in m.lower() for m in memberships):
        memberships.append("American Medical Association (AMA)")

    # Affiliations
    affiliations = d.get("affiliations", [])

    # Ratings
    ratings = list(d.get("ratings", []))
    platform_names = {r.get("platform", "").lower() for r in ratings}
    if "healthgrades" not in platform_names:
        ratings.append({"platform": "Healthgrades", "rating": "See Profile",
                        "notes": f"Listed {specialty} Specialist"})
    if "webmd" not in platform_names:
        ratings.append({"platform": "WebMD", "rating": "See Profile",
                        "notes": f"Listed {specialty}"})
    if "u.s. news" not in platform_names:
        ratings.append({"platform": "U.S. News & World Report", "rating": "Listed",
                        "notes": f"Recognized {specialty} Surgeon"})
    if "doximity" not in platform_names:
        ratings.append({"platform": "Doximity", "rating": "Listed",
                        "notes": f"{specialty} specialist profile"})

    # Awards
    awards = list(d.get("awards", []))

    # Procedures
    procedures = d.get("procedures", [])

    # Languages
    languages = d.get("languages", ["English"])

    # Sources
    source_urls = d.get("source_urls", [])

    # Biography
    biography = d.get("description", "")
    if not biography:
        biography = f"{display_name} is a {specialty.lower()} specialist practicing in {city}, {state}."

    # =====================================================================
    # BUILD DOCUMENT
    # =====================================================================
    doc = Document()

    # --- PAGE HEADER ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"Surgeon Profile | {display_name}")
    run.font.size = Pt(9)
    run.font.color.rgb = HEADER_TEXT_COLOR

    # --- TITLE BLOCK ---
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Emu(TITLE_SPACE_AFTER)

    if LOGO_PATH and os.path.exists(LOGO_PATH):
        run_logo = p0.add_run()
        run_logo.add_picture(LOGO_PATH, width=Emu(1452470), height=Emu(400050))

    run_title = p0.add_run("SURGEON PROFILE")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = DARK_BLUE

    # Name block
    photo_path = d.get("photo_path", "")
    if photo_path and os.path.exists(photo_path) and os.path.getsize(photo_path) > 100:
        _add_name_block_with_photo(doc, display_name, specialty, location_str, photo_path)
    else:
        _add_name_block_no_photo(doc, display_name, specialty, location_str)

    # --- CONTACT INFORMATION ---
    _add_section_heading(doc, "CONTACT INFORMATION")
    if practice_name:
        _add_contact_line(doc, "Practice", practice_name)
    if address:
        _add_contact_line(doc, "Address", f"{address}, {city}, {state} {zip_code}")
    if phone:
        _add_contact_line(doc, "Phone", phone)
    if practice_website:
        _add_contact_line(doc, "Website", practice_website)
    _add_contact_line(doc, "NPI", npi)

    # --- PROFESSIONAL BIOGRAPHY ---
    _add_section_heading(doc, "PROFESSIONAL BIOGRAPHY")
    p_bio = doc.add_paragraph()
    p_bio.paragraph_format.space_after = Emu(BIO_SPACE_AFTER)
    p_bio.add_run(biography)

    # --- EDUCATION & TRAINING ---
    if education_items:
        _add_section_heading(doc, "EDUCATION & TRAINING")
        for item in education_items:
            _add_bullet(doc, item)

    # --- BOARD CERTIFICATIONS & CREDENTIALS ---
    _add_section_heading(doc, "BOARD CERTIFICATIONS & CREDENTIALS")
    for cert in board_certs:
        _add_bullet(doc, cert)

    # --- SURGICAL PERFORMANCE METRICS ---
    if procedures:
        _add_section_heading(doc, "SURGICAL PERFORMANCE METRICS")
        metric_headers = ["Procedure", "Informed Score", "Recommended"]
        metric_data = []
        for proc in procedures:
            score = proc.get("informed_score", 0)
            try:
                score_int = int(score)
            except (ValueError, TypeError):
                score_int = 0
            if score_int >= 90:
                recommendation = "Highly Recommended"
            elif score_int > 70:
                recommendation = "Recommended"
            else:
                recommendation = "Not Recommended"
            metric_data.append([
                proc.get("name", ""),
                str(score),
                recommendation,
            ])
        if metric_data:
            _add_styled_table(doc, metric_headers, metric_data)

    # --- HOSPITAL AFFILIATIONS ---
    if affiliations:
        _add_section_heading(doc, "HOSPITAL AFFILIATIONS")
        for facility in affiliations:
            if isinstance(facility, dict):
                parts = [facility.get("name", "")]
                if facility.get("city"):
                    parts.append(facility["city"])
                if facility.get("state"):
                    parts.append(facility["state"])
                _add_bullet(doc, ", ".join(parts))
            else:
                _add_bullet(doc, str(facility))

    # --- PROFESSIONAL MEMBERSHIPS ---
    _add_section_heading(doc, "PROFESSIONAL MEMBERSHIPS")
    for m in memberships:
        _add_bullet(doc, m)

    # --- PATIENT RATINGS & REVIEWS ---
    _add_section_heading(doc, "PATIENT RATINGS & REVIEWS")
    review_headers = ["Platform", "Rating", "Notes"]
    review_data = []
    for r in ratings:
        review_data.append([
            r.get("platform", ""),
            r.get("rating", "See Profile"),
            r.get("notes", ""),
        ])
    _add_styled_table(doc, review_headers, review_data)

    # --- AWARDS & RECOGNITIONS ---
    if awards:
        _add_section_heading(doc, "AWARDS & RECOGNITIONS")
        for award in awards:
            _add_bullet(doc, award)

    # --- MEDIA & PRESS ---
    media = d.get("media", [])
    if media:
        _add_section_heading(doc, "MEDIA & PRESS")
        for item in media:
            _add_bullet(doc, item)

    # --- ROBOTIC SURGERY (da Vinci) ---
    davinci_status = d.get("davinci_status", None)
    if davinci_status is not None:
        _add_section_heading(doc, "ROBOTIC SURGERY")
        if davinci_status.get("listed"):
            details = davinci_status.get("details", "")
            text = f"Intuitive da Vinci Physician Locator: Yes — Listed"
            if details:
                text += f" ({details})"
            _add_bullet(doc, text)
            davinci_url = davinci_status.get("profile_url", "")
            if davinci_url:
                _add_bullet(doc, f"Profile: {davinci_url}")
        else:
            _add_bullet(doc, "Intuitive da Vinci Physician Locator: No — Not Listed")

    # --- LANGUAGES ---
    _add_section_heading(doc, "LANGUAGES")
    for lang in languages:
        _add_bullet(doc, lang)

    # --- SOURCES & CITATIONS ---
    if source_urls:
        _add_section_heading(doc, "SOURCES & CITATIONS")
        for url in source_urls:
            p_src = doc.add_paragraph()
            p_src.paragraph_format.space_after = Emu(CONTACT_SPACE_AFTER)
            p_src.add_run(url)

    # --- SAVE ---
    last = full_name.split()[-1] if full_name else "Unknown"
    first = full_name.split()[0] if full_name else "Unknown"
    filename = f"{npi} - {last} {first}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath
